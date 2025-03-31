import whisper
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
import torch
import threading

class TranscriptionService:
    def __init__(self, app=None):
        # Check if CUDA is available for GPU acceleration
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model = None
        self.model_lock = threading.Lock()
        self.app = app
        
        if app:
            app.logger.info(f"Transcription service initialized (model will be loaded on demand)")
    
    def _log(self, message, level="info"):
        """Log a message using the app logger if available"""
        if self.app:
            logger = getattr(self.app.logger, level)
            logger(message)
        else:
            print(f"{level.upper()}: {message}")
    
    def _load_model(self):
        """Load the Whisper model if not already loaded"""
        with self.model_lock:
            if self.model is None:
                try:
                    self._log(f"Loading Whisper model on {self.device}...")
                    self.model = whisper.load_model("medium", device=self.device)
                    self._log(f"Whisper model loaded successfully on {self.device}")
                except Exception as e:
                    self._log(f"Failed to load Whisper model: {str(e)}", "error")
                    raise
    
    def transcribe_audio(self, audio_path, language='en'):
        """Transcribe audio file using Whisper"""
        try:
            # Log the transcription start
            self._log(f"Starting transcription of {audio_path}")
            
            # Ensure the file exists
            if not os.path.exists(audio_path):
                raise FileNotFoundError(f"Audio file not found: {audio_path}")
            
            # Load model if not already loaded
            if self.model is None:
                self._load_model()
            
            # Perform transcription
            result = self.model.transcribe(
                audio_path,
                language=language,
                fp16=torch.cuda.is_available()  # Use FP16 if on GPU
            )
            
            self._log(f"Transcription completed for {audio_path}")
            return result['text'], None
            
        except Exception as e:
            self._log(f"Transcription error: {str(e)}", "error")
            return None, str(e)
    
    def create_transcript_document(self, transcription_text, metadata):
        """Create a formatted Word document with the transcription"""
        try:
            doc = Document()
            
            # Set document metadata
            doc.core_properties.title = metadata.get('title', 'Video Transcript')
            doc.core_properties.author = metadata.get('author', '')
            doc.core_properties.keywords = metadata.get('keywords', '')
            doc.core_properties.comments = metadata.get('comments', '')
            
            # Add headers
            headers = [
                f"{metadata.get('platform', '')} VIDEO TRANSCRIPT",
                f"{metadata.get('course_code', '')} – {metadata.get('course_name', '')}",
                f"WEEK {metadata.get('week_no', '')} {metadata.get('transcript_type', '')} {metadata.get('part', '')} Transcript",
                metadata.get('week_topic', '')
            ]
            
            for header in headers:
                p = doc.add_paragraph()
                run = p.add_run(header)
                p.style = 'Heading 1'
                run.font.size = Pt(18)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_before = Pt(0)
                p.space_after = Pt(0)
            
            # Add transcription paragraphs
            paragraphs = self._split_text_into_paragraphs(transcription_text)
            for paragraph in paragraphs:
                p = doc.add_paragraph(paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(16)
            
            return doc, None
            
        except Exception as e:
            self._log(f"Document creation error: {str(e)}", "error")
            return None, str(e)
    
    def _split_text_into_paragraphs(self, text, max_sentences=5):
        """Split text into paragraphs based on sentence boundaries"""
        sentences = re.split(r'(?<=[.!?]) +', text)
        paragraphs = []
        current_paragraph = []
        
        for sentence in sentences:
            current_paragraph.append(sentence)
            if len(current_paragraph) >= max_sentences:
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []
        
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        return paragraphs 